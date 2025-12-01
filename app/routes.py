import os
import secrets
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import (render_template, flash, redirect, url_for, request,
                   Blueprint, abort, current_app, jsonify)
from flask_login import current_user, login_user, logout_user, login_required
from app import db
from app.models import User, IpAsset, Contract, IpAnalytics
from app.forms import LoginForm, RegistrationForm, IpAssetForm, ContractForm
from sqlalchemy import func, desc, or_
from functools import wraps

bp = Blueprint('main', __name__)


# --- 辅助函数 ---
def save_picture(form_picture):
    random_hex = secrets.token_hex(8)
    _, f_ext = os.path.splitext(form_picture.filename)
    picture_fn = random_hex + f_ext
    picture_path = os.path.join(current_app.root_path, 'static/images', picture_fn)
    os.makedirs(os.path.dirname(picture_path), exist_ok=True)
    form_picture.save(picture_path)
    return os.path.join('images', picture_fn)


def save_pdf(form_pdf):
    random_hex = secrets.token_hex(8)
    _, f_ext = os.path.splitext(form_pdf.filename)
    pdf_fn = random_hex + f_ext
    pdf_path = os.path.join(current_app.root_path, 'static/pdfs', pdf_fn)
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    form_pdf.save(pdf_path)
    return os.path.join('pdfs', pdf_fn)


def get_licensing_status(ip_asset):
    today = date.today()
    active_contracts = Contract.query.filter(
        Contract.ip_id == ip_asset.id,
        Contract.license_type.in_(['独占许可', '排他许可']),
        Contract.term_start <= today,
        Contract.term_end >= today
    ).all()
    if not active_contracts: return "✅ 可全球许可"
    regions = [c.region for c in active_contracts if c.region]
    for r in regions:
        if "全球" in r: return "🔒 不可许可(全球独占)"
    if regions:
        region_str = "、".join(list(set(regions)))
        return f"⚠️ 除{region_str}外可许可"
    return "✅ 可许可"


# --- 权限装饰器 ---
def internal_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or current_user.role != 'internal': abort(403)
        return f(*args, **kwargs)

    return decorated_function


def partner_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or current_user.role != 'partner': abort(403)
        return f(*args, **kwargs)

    return decorated_function


# --- 页面路由 ---
@bp.route('/')
@bp.route('/index')
@login_required
def index():
    if current_user.role == 'internal':
        return redirect(url_for('main.internal_dashboard'))
    elif current_user.role == 'partner':
        return redirect(url_for('main.portal_dashboard'))
    else:
        return redirect(url_for('main.login'))


@bp.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated: return redirect(url_for('main.index'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user is None or not user.check_password(form.password.data):
            flash('无效用户名或密码', 'danger');
            return redirect(url_for('main.login'))
        login_user(user, remember=form.remember_me.data)
        return redirect(url_for('main.index'))
    return render_template('login.html', title='登录', form=form)


@bp.route('/logout')
def logout(): logout_user(); return redirect(url_for('main.login'))


@bp.route('/register', methods=['GET', 'POST'])
def register():
    form = RegistrationForm()
    if form.validate_on_submit():
        user = User(username=form.username.data, role=form.role.data)
        user.set_password(form.password.data)
        db.session.add(user);
        db.session.commit()
        return redirect(url_for('main.login'))
    return render_template('register.html', title='注册', form=form)


# --- 内控端页面 ---
@bp.route('/internal/dashboard')
@login_required
@internal_required
def internal_dashboard():
    rev_stats = db.session.query(IpAsset.name, IpAsset.current_revenue).order_by(desc(IpAsset.current_revenue)).limit(
        5).all()
    rev_x = [r[0] for r in rev_stats];
    rev_y = [float(r[1]) for r in rev_stats]
    click_stats = db.session.query(IpAsset.name, func.count(IpAnalytics.id)).join(IpAnalytics).group_by(
        IpAsset.name).order_by(func.count(IpAnalytics.id).desc()).limit(5).all()
    click_x = [c[0] for c in click_stats];
    click_y = [c[1] for c in click_stats]

    ip_q_name = request.args.get('ip_name', '');
    ip_q_author = request.args.get('ip_author', '')
    ip_query = IpAsset.query
    if ip_q_name: ip_query = ip_query.filter(IpAsset.name.like(f'%{ip_q_name}%'))
    if ip_q_author: ip_query = ip_query.filter(IpAsset.author.like(f'%{ip_q_author}%'))
    ips = ip_query.all()
    ip_statuses = {ip.id: get_licensing_status(ip) for ip in ips}

    ct_q_partner = request.args.get('ct_partner', '');
    ct_q_ip = request.args.get('ct_ip', '');
    ct_q_type = request.args.get('ct_type', '')
    ct_query = Contract.query.join(IpAsset)
    if ct_q_partner: ct_query = ct_query.filter(Contract.partner_name.like(f'%{ct_q_partner}%'))
    if ct_q_ip: ct_query = ct_query.filter(IpAsset.name.like(f'%{ct_q_ip}%'))
    if ct_q_type: ct_query = ct_query.filter(Contract.license_type == ct_q_type)
    contracts = ct_query.order_by(desc(Contract.id)).all()

    ip_form = IpAssetForm();
    contract_form = ContractForm()
    contract_form.ip_id.choices = [(i.id, i.name) for i in IpAsset.query.all()]
    tencent_embed_url = os.environ.get('TENCENT_EMBED_URL_INTERNAL')

    return render_template('internal_dashboard.html', title='内控台账', ips=ips, ip_statuses=ip_statuses,
                           contracts=contracts, rev_x=rev_x, rev_y=rev_y, click_x=click_x, click_y=click_y,
                           ip_form=ip_form, contract_form=contract_form, search_ip_name=ip_q_name,
                           search_ip_author=ip_q_author, search_ct_partner=ct_q_partner, search_ct_ip=ct_q_ip,
                           search_ct_type=ct_q_type, tencent_embed_url=tencent_embed_url)


# --- 伙伴端页面 ---
@bp.route('/portal/dashboard')
@login_required
@partner_required
def portal_dashboard():
    query = request.args.get('q', '')
    base_query = IpAsset.query
    if query: base_query = base_query.filter(or_(IpAsset.name.like(f'%{query}%'), IpAsset.tags.like(f'%{query}%')))
    ips = base_query.all()
    cases = Contract.query.filter(Contract.case_image_url != None).order_by(desc(Contract.id)).limit(8).all()
    tencent_embed_url = os.environ.get('TENCENT_EMBED_URL_PARTNER')
    return render_template('portal_dashboard.html', title='合作伙伴端', ips=ips, cases=cases, search_query=query,
                           tencent_embed_url=tencent_embed_url)


@bp.route('/portal/ip/<int:ip_id>')
@login_required
@partner_required
def ip_detail(ip_id):
    ip = IpAsset.query.get_or_404(ip_id)
    try:
        db.session.add(IpAnalytics(ip_id=ip.id, user_id=current_user.id)); db.session.commit()
    except:
        pass
    return render_template('ip_detail.html', ip=ip)


# --- 增删操作 ---
@bp.route('/ip/add', methods=['POST'])
@login_required
@internal_required
def add_ip():
    form = IpAssetForm()
    if form.validate_on_submit():
        img = save_picture(form.image_file.data) if form.image_file.data else None
        new_ip = IpAsset(name=form.name.data, tags=form.tags.data, description=form.description.data, image_url=img,
                         author=form.author.data, ownership=form.ownership.data, reg_number=form.reg_number.data,
                         reg_date=form.reg_date.data, trademark_info=form.trademark_info.data,
                         license_period=form.license_period.data, contact_email=form.contact_email.data,
                         license_type_options=form.license_type_options.data, value_level=form.value_level.data,
                         current_revenue=form.current_revenue.data or 0)
        db.session.add(new_ip);
        db.session.commit();
        flash('IP 添加成功', 'success')
    else:
        flash('添加失败', 'danger')
    return redirect(url_for('main.internal_dashboard'))


@bp.route('/ip/delete/<int:ip_id>', methods=['POST'])
@login_required
@internal_required
def delete_ip(ip_id):
    ip = IpAsset.query.get_or_404(ip_id)
    if ip.contracts:
        flash('无法删除：请先删除关联的合同', 'danger')
    else:
        db.session.delete(ip); db.session.commit(); flash('IP 已删除', 'success')
    return redirect(url_for('main.internal_dashboard'))


@bp.route('/contract/add', methods=['POST'])
@login_required
@internal_required
def add_contract():
    form = ContractForm()
    form.ip_id.choices = [(i.id, i.name) for i in IpAsset.query.all()]
    if form.validate_on_submit():
        img = save_picture(form.case_image_file.data) if form.case_image_file.data else None
        pdf = save_pdf(form.pdf_file.data) if form.pdf_file.data else None
        nc = Contract(ip_id=form.ip_id.data, partner_name=form.partner_name.data, partner_brand=form.partner_brand.data,
                      region=form.region.data, media=form.media.data, license_method=form.license_method.data,
                      license_category=form.license_category.data, usage_type=form.usage_type.data,
                      license_type=form.license_type.data, term_start=form.term_start.data, term_end=form.term_end.data,
                      fee_standard=form.fee_standard.data, payment_cycle=form.payment_cycle.data,
                      breach_terms=form.breach_terms.data, case_image_url=img, pdf_url=pdf)
        db.session.add(nc);
        db.session.commit();
        flash('合同添加成功', 'success')
    else:
        flash('添加失败', 'danger')
    return redirect(url_for('main.internal_dashboard'))


@bp.route('/contract/delete/<int:contract_id>', methods=['POST'])
@login_required
@internal_required
def delete_contract(contract_id):
    c = Contract.query.get_or_404(contract_id)
    db.session.delete(c);
    db.session.commit();
    flash('合同已删除', 'success')
    return redirect(url_for('main.internal_dashboard'))


# ==================================================================
# --- 伙伴端专用 API ---
# ==================================================================

@bp.route('/api/partner/get_licensable_ips', methods=['POST'])
def partner_get_licensable_ips():
    try:
        ips = IpAsset.query.all()
        # 优化：只返回关键信息，减少 Token 消耗
        report = []
        for ip in ips:
            info = f"名称:{ip.name}|标签:{ip.tags}|描述:{ip.description[:50]}...|级别:{ip.value_level}"
            report.append(info)
        return jsonify({"status": "success", "count": len(ips), "ips_report": "\n".join(report)})
    except Exception as e:
        print(f"Error Partner IP: {e}")  # 打印日志
        return jsonify({"status": "error", "message": str(e)}), 500


@bp.route('/api/partner/get_fee_guidance', methods=['POST'])
def partner_get_fee_guidance():
    data = request.get_json() or {}
    ip_name = data.get('ip_name')
    if not ip_name: return jsonify({"status": "error", "message": "Missing ip_name"}), 400
    try:
        ip = IpAsset.query.filter(IpAsset.name == ip_name).first()
        if not ip: return jsonify({"status": "error", "message": "IP not found", "value_level": "未知"})
        return jsonify({"status": "success", "ip_name": ip.name, "value_level": ip.value_level})
    except Exception as e:
        print(f"Error Partner Fee: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500


# ==================================================================
# --- 内控端专用 API (修复版) ---
# ==================================================================

# app/routes.py 中找到对应的部分进行替换

@bp.route('/api/internal/get_database_info', methods=['GET', 'POST'])
def internal_get_database_info():
    """
    [内控端] 模块1: 上帝视角全量查询 (修复版: 支持GET测试，防超时)
    """
    try:
        # 1. IP 资产 (只取前20条，防止超时)
        ips = IpAsset.query.limit(20).all()
        report = ["【IP资产(前20条)】"]
        for ip in ips:
            # 简化状态描述，减少计算量
            status = ip.internal_status
            line = (f"ID:{ip.id}|名:{ip.name}|级:{ip.value_level}|益:{ip.current_revenue}|"
                    f"权:{ip.ownership}|商:{ip.trademark_info}|态:{status}")
            report.append(line)

        # 2. 合同数据 (只取前20条，防止超时)
        contracts = Contract.query.order_by(desc(Contract.id)).limit(20).all()
        for c in contracts:
            # 截取过长的文本
            term_info = f"{c.term_start}~{c.term_end}"
            line = (f"ID:{c.id}|方:{c.partner_name}|IP:{c.ip_asset.name}|型:{c.license_type}|"
                    f"期:{term_info}|费:{c.fee_standard[:10]}...")
            report.append(line)

        final_report = "\n".join(report)

        print(f"API Success. Length: {len(final_report)}")  # 打印成功日志
        return jsonify({"status": "success", "info_report": final_report})

    except Exception as e:
        # 将错误打印到 PythonAnywhere 的 Server Log
        import traceback
        traceback.print_exc()
        print(f"API ERROR: {str(e)}")
        return jsonify({"status": "error", "message": f"Server Error: {str(e)}"}), 500


@bp.route('/api/internal/get_report', methods=['POST'])
def internal_get_report():
    try:
        rev_stats = db.session.query(IpAsset.name, IpAsset.current_revenue).order_by(
            desc(IpAsset.current_revenue)).limit(5).all()
        rev_rpt = "【收益Top5】\n" + "\n".join([f"{r[0]}: {r[1]}万" for r in rev_stats])

        click_stats = db.session.query(IpAsset.name, func.count(IpAnalytics.id)).join(IpAnalytics).group_by(
            IpAsset.name).order_by(func.count(IpAnalytics.id).desc()).limit(5).all()
        click_rpt = "【热度Top5】\n" + "\n".join([f"{c[0]}: {c[1]}次" for c in click_stats])

        today = date.today();
        ninety_days = today + timedelta(days=90)
        expiring = Contract.query.filter(Contract.term_end >= today, Contract.term_end <= ninety_days).order_by(
            Contract.term_end).all()
        exp_rpt = "【到期预警】\n" + (
            "\n".join([f"{c.partner_name}({c.ip_asset.name}):{c.term_end}" for c in expiring]) if expiring else "无")

        return jsonify(
            {"status": "success", "revenue_report": rev_rpt, "click_report": click_rpt, "expire_report": exp_rpt})
    except Exception as e:
        print(f"Error Report: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500


@bp.route('/api/internal/generate_contract_doc', methods=['POST'])
def internal_generate_contract_doc():
    """
    专业版合同生成：包含标准法务条款
    """
    data = request.get_json() or {}
    try:
        doc = Document()

        # 设置标题样式
        title = doc.add_heading('IP 授权许可合同', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 提取参数
        p_name = data.get('partner_name', '______')
        ip_name = data.get('ip_name', '______')
        region = data.get('region', '中国大陆')
        media = data.get('media', '全媒体')
        usage = data.get('usage_type', '商业授权')
        c_type = data.get('license_type', '普通许可')
        term = data.get('term', '1年')
        start_d = data.get('start_date', str(date.today()))
        fee = data.get('fee', '待定')
        cycle = data.get('payment_cycle', '一次性支付')
        breach = data.get('breach_terms', '依法承担赔偿责任')

        # 头部
        doc.add_paragraph(f'合同编号：IP-{secrets.token_hex(3).upper()}-{date.today().year}')
        doc.add_paragraph(f'\n甲方（授权方）：星核文化科技发展有限公司')
        doc.add_paragraph(f'乙方（被授权方）：{p_name}')
        doc.add_paragraph(f'签署日期：{date.today().strftime("%Y年%m月%d日")}')

        doc.add_paragraph(
            '\n鉴于甲方拥有本合同所述IP的合法著作权，乙方希望获得该IP的使用授权，双方经友好协商，达成如下协议：')

        # 第一条
        h1 = doc.add_heading('第一条 授权内容', level=1)
        doc.add_paragraph(f'1.1 授权标的：甲方授权乙方使用 IP “{ip_name}” 的形象及相关素材。')
        doc.add_paragraph(f'1.2 授权地域：{region}。')
        doc.add_paragraph(f'1.3 授权渠道/媒介：{media}。')
        doc.add_paragraph(f'1.4 具体使用方式：{usage}。')
        doc.add_paragraph(f'1.5 许可性质：本授权为【{c_type}】。')

        # 第二条
        doc.add_heading('第二条 授权期限', level=1)
        doc.add_paragraph(f'2.1 本合同有效期为 {term}，自 {start_d} 起生效。')
        doc.add_paragraph('2.2 期限届满前30日，双方可协商续约事宜。')

        # 第三条
        doc.add_heading('第三条 费用与支付', level=1)
        doc.add_paragraph(f'3.1 授权费用：总金额为人民币 {fee}。')
        doc.add_paragraph(f'3.2 支付方式：{cycle}。乙方应将款项汇入甲方指定账户。')
        doc.add_paragraph('3.3 税费承担：除另有约定外，相关税费由各自依法承担。')

        # 第四条
        doc.add_heading('第四条 知识产权声明', level=1)
        doc.add_paragraph('4.1 甲方保证对授权IP享有完整的知识产权，未侵犯任何第三方的合法权益。')
        doc.add_paragraph(
            '4.2 乙方在使用过程中产生的新设计（衍生品设计图等），其知识产权归属双方另行约定；若未约定，原则上归甲方所有。')

        # 第五条
        doc.add_heading('第五条 保密义务', level=1)
        doc.add_paragraph('5.1 双方对本合同内容及在合作过程中获知的对方商业秘密负有保密义务，不得向第三方披露。')

        # 第六条 (动态填入违约责任)
        doc.add_heading('第六条 违约责任', level=1)
        doc.add_paragraph(f'6.1 {breach}')
        doc.add_paragraph('6.2 若乙方超出授权范围使用IP，甲方有权单方解除合同并要求赔偿。')

        # 第七条
        doc.add_heading('第七条 争议解决', level=1)
        doc.add_paragraph('7.1 因本合同引起的任何争议，双方应友好协商解决；协商不成的，应向甲方所在地人民法院提起诉讼。')

        # 签字栏
        doc.add_paragraph('\n\n（以下无正文）\n')
        table = doc.add_table(rows=1, cols=2)
        row = table.rows[0]
        row.cells[0].text = "甲方：星核文化科技发展有限公司\n\n代表签字：__________________"
        row.cells[1].text = f"乙方：{p_name}\n\n代表签字：__________________"

        # 保存
        fn = f"Contract_Pro_{secrets.token_hex(4)}.docx"
        sd = os.path.join(current_app.root_path, 'static', 'generated_docs')
        if not os.path.exists(sd): os.makedirs(sd)
        doc.save(os.path.join(sd, fn))

        return jsonify({
            "status": "success",
            "download_url": url_for('static', filename=f'generated_docs/{fn}', _external=True)
        })
    except Exception as e:
        print(f"Gen Doc Error: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500